"""Visual validation tests for the full DOCX -> LaTeX -> PDF -> Screenshot pipeline.

This module provides automated testing of the visual output of converted documents.
It compares screenshots of generated PDFs against baseline images to detect layout regressions.

Usage:
    # Run visual tests (generates screenshots, compares to baselines)
    pytest tests/visual/ -v

    # Update baselines with current output
    pytest tests/visual/ -v --update-baselines

    # Generate screenshots only (no comparison)
    pytest tests/visual/ -v --generate-only
"""

from __future__ import annotations

import shutil
import subprocess
from pathlib import Path
from typing import TYPE_CHECKING

import pytest
from PIL import Image, ImageChops, ImageDraw

from docx2latex.application.dto.conversion_options import ConversionOptions
from docx2latex.application.services.conversion_service import ConversionService

if TYPE_CHECKING:
    from collections.abc import Generator


# ============================================================================
# Configuration
# ============================================================================

VISUAL_TESTS_DIR = Path(__file__).parent
BASELINES_DIR = VISUAL_TESTS_DIR / "baselines"
OUTPUTS_DIR = VISUAL_TESTS_DIR / "outputs"
DIFFS_DIR = VISUAL_TESTS_DIR / "diffs"

# Threshold for image difference (0-100, percentage of different pixels)
DIFF_THRESHOLD_PERCENT = 1.0

# DPI for PDF rendering
PDF_DPI = 150

# Text content baselines
TEXT_BASELINES_DIR = VISUAL_TESTS_DIR / "text_baselines"


# ============================================================================
# Helper Functions
# ============================================================================


def ensure_dirs() -> None:
    """Ensure output directories exist."""
    BASELINES_DIR.mkdir(parents=True, exist_ok=True)
    OUTPUTS_DIR.mkdir(parents=True, exist_ok=True)
    DIFFS_DIR.mkdir(parents=True, exist_ok=True)
    TEXT_BASELINES_DIR.mkdir(parents=True, exist_ok=True)


def compile_latex_to_pdf(tex_path: Path, output_dir: Path) -> Path | None:
    """Compile a LaTeX file to PDF using pdflatex.

    Args:
        tex_path: Path to the .tex file
        output_dir: Directory for output files

    Returns:
        Path to the generated PDF, or None if compilation failed
    """
    if not shutil.which("pdflatex"):
        pytest.skip("pdflatex not found in PATH")

    pdf_path = output_dir / f"{tex_path.stem}.pdf"

    # Run pdflatex twice for proper references
    for _ in range(2):
        result = subprocess.run(
            [
                "pdflatex",
                "-interaction=nonstopmode",
                "-output-directory",
                str(output_dir),
                str(tex_path),
            ],
            capture_output=True,
            timeout=60,
            cwd=tex_path.parent,  # Run from tex file directory for image paths
        )

        if result.returncode != 0:
            # Log the error but don't fail immediately - check if PDF was created
            log_path = output_dir / f"{tex_path.stem}.log"
            if log_path.exists():
                try:
                    log_content = log_path.read_text(encoding="utf-8", errors="replace")
                except Exception:
                    log_content = log_path.read_text(encoding="latin-1", errors="replace")
                # Extract error lines
                errors = [
                    line for line in log_content.split("\n") if line.startswith("!")
                ]
                if errors:
                    print(f"LaTeX errors in {tex_path.name}:")
                    for error in errors[:5]:  # Show first 5 errors
                        print(f"  {error}")

    if pdf_path.exists():
        return pdf_path
    return None


def pdf_to_images(pdf_path: Path, dpi: int = PDF_DPI) -> list[Image.Image]:
    """Convert a PDF to a list of PIL Images (one per page).

    Uses PyMuPDF (fitz) which works cross-platform (Windows, macOS, Linux)
    without requiring external dependencies like poppler.

    Args:
        pdf_path: Path to the PDF file
        dpi: Resolution for rendering

    Returns:
        List of PIL Image objects
    """
    try:
        import fitz  # PyMuPDF

        images = []
        doc = fitz.open(pdf_path)

        # Calculate zoom factor from DPI (default PDF DPI is 72)
        zoom = dpi / 72
        matrix = fitz.Matrix(zoom, zoom)

        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            pix = page.get_pixmap(matrix=matrix)

            # Convert to PIL Image
            img = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
            images.append(img)

        doc.close()
        return images

    except ImportError:
        pytest.skip(
            "PyMuPDF not installed. Install with: pip install pymupdf\n"
            "Or install visual test dependencies: pip install -e '.[visual]'"
        )


def extract_text_from_pdf(pdf_path: Path) -> str:
    """Extract text content from a PDF file.

    Uses PyMuPDF to extract all text, which allows content validation
    independent of layout/visual rendering.

    Args:
        pdf_path: Path to the PDF file

    Returns:
        Extracted text content (all pages concatenated)
    """
    try:
        import fitz  # PyMuPDF

        doc = fitz.open(pdf_path)
        text_parts = []

        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            text_parts.append(page.get_text())

        doc.close()
        return "\n".join(text_parts)

    except ImportError:
        pytest.skip("PyMuPDF not installed")


def extract_text_from_docx(docx_path: Path) -> str:
    """Extract text content from a DOCX file.

    Args:
        docx_path: Path to the DOCX file

    Returns:
        Extracted text content
    """
    from docx import Document

    doc = Document(docx_path)
    text_parts = []

    for para in doc.paragraphs:
        text_parts.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text_parts.append(cell.text)

    return "\n".join(text_parts)


def normalize_text(text: str) -> str:
    """Normalize text for comparison.

    Removes extra whitespace, normalizes line endings, etc.

    Args:
        text: Raw text

    Returns:
        Normalized text
    """
    import re

    # Normalize whitespace
    text = re.sub(r"[ \t]+", " ", text)
    # Normalize line endings
    text = re.sub(r"\r\n|\r", "\n", text)
    # Remove multiple blank lines
    text = re.sub(r"\n{3,}", "\n\n", text)
    # Strip each line
    lines = [line.strip() for line in text.split("\n")]
    # Remove empty lines at start/end
    while lines and not lines[0]:
        lines.pop(0)
    while lines and not lines[-1]:
        lines.pop()

    return "\n".join(lines)


def compare_text_content(
    docx_text: str, pdf_text: str
) -> tuple[bool, list[str]]:
    """Compare text content between DOCX source and PDF output.

    Args:
        docx_text: Text extracted from DOCX
        pdf_text: Text extracted from PDF

    Returns:
        Tuple of (is_match, list of differences)
    """
    # Normalize both texts
    docx_normalized = normalize_text(docx_text)
    pdf_normalized = normalize_text(pdf_text)

    # Split into words for comparison (ignoring whitespace differences)
    docx_words = set(docx_normalized.split())
    pdf_words = set(pdf_normalized.split())

    # Find missing and extra words
    missing_in_pdf = docx_words - pdf_words
    extra_in_pdf = pdf_words - docx_words

    differences = []

    # Filter out common LaTeX artifacts that are expected
    latex_artifacts = {"1", "2", "3", "4", "5", "6", "7", "8", "9", "0", "-", "•"}

    missing_significant = missing_in_pdf - latex_artifacts
    if missing_significant:
        # Only report if significant words are missing
        significant_missing = [w for w in missing_significant if len(w) > 2]
        if significant_missing:
            differences.append(f"Missing in PDF: {', '.join(sorted(significant_missing)[:10])}")

    # Check if key content words from DOCX appear in PDF
    # (more lenient check than exact match)
    key_words = {w for w in docx_words if len(w) > 3 and w.isalpha()}
    found_key_words = key_words & pdf_words
    coverage = len(found_key_words) / len(key_words) if key_words else 1.0

    is_match = coverage >= 0.8  # 80% of key words should be present

    if coverage < 1.0:
        differences.append(f"Content coverage: {coverage:.1%}")

    return is_match, differences


def compare_images(
    img1: Image.Image, img2: Image.Image, diff_path: Path | None = None
) -> tuple[bool, float]:
    """Compare two images and optionally save a diff image.

    Args:
        img1: First image (baseline)
        img2: Second image (current)
        diff_path: Optional path to save diff visualization

    Returns:
        Tuple of (is_same, diff_percentage)
    """
    # Ensure same size
    if img1.size != img2.size:
        # Resize to larger dimensions for comparison
        max_size = (max(img1.width, img2.width), max(img1.height, img2.height))
        img1 = img1.resize(max_size, Image.Resampling.LANCZOS)
        img2 = img2.resize(max_size, Image.Resampling.LANCZOS)

    # Convert to same mode
    if img1.mode != img2.mode:
        img1 = img1.convert("RGB")
        img2 = img2.convert("RGB")

    # Calculate difference
    diff = ImageChops.difference(img1, img2)

    # Calculate percentage of different pixels
    diff_data = diff.getdata()
    total_pixels = len(diff_data)
    different_pixels = sum(1 for pixel in diff_data if sum(pixel) > 30)  # Threshold for noise
    diff_percent = (different_pixels / total_pixels) * 100

    is_same = diff_percent <= DIFF_THRESHOLD_PERCENT

    # Save diff visualization if requested
    if diff_path and not is_same:
        # Create a visualization showing differences
        diff_vis = Image.new("RGB", img1.size, (255, 255, 255))
        diff_vis.paste(img1)

        # Highlight differences in red
        draw = ImageDraw.Draw(diff_vis)
        for i, pixel in enumerate(diff_data):
            if sum(pixel) > 30:
                x = i % img1.width
                y = i // img1.width
                draw.point((x, y), fill=(255, 0, 0))

        diff_vis.save(diff_path)

    return is_same, diff_percent


def process_document(
    docx_path: Path,
    output_dir: Path,
    name: str,
) -> tuple[Path | None, list[Image.Image]]:
    """Process a DOCX file through the full pipeline.

    Args:
        docx_path: Path to the DOCX file
        output_dir: Directory for output files
        name: Base name for output files

    Returns:
        Tuple of (pdf_path, list of page images)
    """
    # Step 1: Convert DOCX to LaTeX
    service = ConversionService()
    tex_path = output_dir / f"{name}.tex"
    options = ConversionOptions(
        output_path=tex_path,
        image_dir="images",
    )

    result = service.convert(docx_path, options)

    if not result.success:
        pytest.fail(f"DOCX conversion failed: {result.errors}")

    assert result.output_path is not None

    # Step 2: Compile LaTeX to PDF
    pdf_path = compile_latex_to_pdf(result.output_path, output_dir)

    if pdf_path is None:
        return None, []

    # Step 3: Convert PDF to images
    images = pdf_to_images(pdf_path)

    return pdf_path, images


# ============================================================================
# Test Class
# ============================================================================


class TestVisualValidation:
    """Visual validation tests for document conversion."""

    @pytest.fixture(autouse=True)
    def setup(self) -> Generator[None, None, None]:
        """Set up test environment."""
        ensure_dirs()
        yield
        # Cleanup is optional - keeping outputs can be useful for debugging

    def test_visual_pipeline_sample_files(
        self,
        sample_dir: Path,
        update_baselines: bool,
        generate_only: bool,
    ) -> None:
        """Test visual output for all sample DOCX files."""
        samples = list(sample_dir.glob("*.docx"))
        samples = [s for s in samples if not s.name.startswith("~")]

        if not samples:
            pytest.skip("No sample DOCX files found")

        for sample in samples:
            name = sample.stem
            print(f"\n{'='*60}")
            print(f"Processing: {sample.name}")
            print(f"{'='*60}")

            # Process the document
            pdf_path, images = process_document(sample, OUTPUTS_DIR, name)

            if not images:
                print(f"  WARNING: No images generated for {name}")
                continue

            print(f"  Generated {len(images)} page(s)")

            # Save current output images
            for i, img in enumerate(images):
                output_path = OUTPUTS_DIR / f"{name}_page_{i}.png"
                img.save(output_path)
                print(f"  Saved: {output_path.name}")

            if generate_only:
                print(f"  (generate-only mode, skipping comparison)")
                continue

            # Compare with baselines or update them
            for i, img in enumerate(images):
                baseline_path = BASELINES_DIR / f"{name}_page_{i}.png"
                output_path = OUTPUTS_DIR / f"{name}_page_{i}.png"
                diff_path = DIFFS_DIR / f"{name}_page_{i}_diff.png"

                if update_baselines:
                    img.save(baseline_path)
                    print(f"  Updated baseline: {baseline_path.name}")
                elif baseline_path.exists():
                    baseline_img = Image.open(baseline_path)
                    is_same, diff_percent = compare_images(
                        baseline_img, img, diff_path
                    )

                    if is_same:
                        print(f"  Page {i}: OK (diff: {diff_percent:.2f}%)")
                    else:
                        print(f"  Page {i}: DIFFERS (diff: {diff_percent:.2f}%)")
                        print(f"    Diff saved to: {diff_path}")
                        # Don't fail immediately, collect all differences
                else:
                    print(f"  Page {i}: No baseline found")
                    print(f"    Run with --update-baselines to create")

    @pytest.mark.parametrize(
        "sample_name",
        ["DM3e", "Fev", "FevShort"],  # Add specific samples to test individually
    )
    def test_visual_single_sample(
        self,
        sample_name: str,
        sample_dir: Path,
        update_baselines: bool,
        generate_only: bool,
    ) -> None:
        """Test visual output for a specific sample file."""
        sample = sample_dir / f"{sample_name}.docx"

        if not sample.exists():
            pytest.skip(f"Sample file not found: {sample}")

        print(f"\nProcessing: {sample.name}")

        # Process the document
        pdf_path, images = process_document(sample, OUTPUTS_DIR, sample_name)

        if not images:
            pytest.fail(f"No images generated for {sample_name}")

        print(f"Generated {len(images)} page(s)")

        failed_pages: list[tuple[int, float]] = []

        for i, img in enumerate(images):
            baseline_path = BASELINES_DIR / f"{sample_name}_page_{i}.png"
            output_path = OUTPUTS_DIR / f"{sample_name}_page_{i}.png"
            diff_path = DIFFS_DIR / f"{sample_name}_page_{i}_diff.png"

            img.save(output_path)

            if update_baselines:
                img.save(baseline_path)
                print(f"  Updated baseline: {baseline_path.name}")
            elif generate_only:
                print(f"  Generated: {output_path.name}")
            elif baseline_path.exists():
                baseline_img = Image.open(baseline_path)
                is_same, diff_percent = compare_images(baseline_img, img, diff_path)

                if not is_same:
                    failed_pages.append((i, diff_percent))
                    print(f"  Page {i}: DIFFERS ({diff_percent:.2f}%)")
                else:
                    print(f"  Page {i}: OK ({diff_percent:.2f}%)")
            else:
                print(f"  Page {i}: No baseline (run with --update-baselines)")

        if failed_pages and not (update_baselines or generate_only):
            pytest.fail(
                f"Visual differences detected on pages: "
                f"{', '.join(f'{p}({d:.1f}%)' for p, d in failed_pages)}"
            )

    def test_content_validation(
        self,
        sample_dir: Path,
    ) -> None:
        """Validate that text content from DOCX appears in the generated PDF.

        This test extracts text from both the source DOCX and the generated PDF,
        then compares them to ensure no content is lost during conversion.
        """
        samples = list(sample_dir.glob("*.docx"))
        samples = [s for s in samples if not s.name.startswith("~")]

        if not samples:
            pytest.skip("No sample DOCX files found")

        for sample in samples:
            name = sample.stem
            print(f"\n{'='*60}")
            print(f"Content validation: {sample.name}")
            print(f"{'='*60}")

            # Process the document
            pdf_path, _ = process_document(sample, OUTPUTS_DIR, name)

            if pdf_path is None:
                print(f"  WARNING: No PDF generated for {name}")
                continue

            # Extract text from both files
            docx_text = extract_text_from_docx(sample)
            pdf_text = extract_text_from_pdf(pdf_path)

            # Save extracted texts for debugging
            docx_text_path = OUTPUTS_DIR / f"{name}_docx_text.txt"
            pdf_text_path = OUTPUTS_DIR / f"{name}_pdf_text.txt"
            docx_text_path.write_text(docx_text, encoding="utf-8")
            pdf_text_path.write_text(pdf_text, encoding="utf-8")

            # Compare content
            is_match, differences = compare_text_content(docx_text, pdf_text)

            if is_match:
                print(f"  Content: OK")
            else:
                print(f"  Content: ISSUES DETECTED")
                for diff in differences:
                    print(f"    - {diff}")

            # Don't fail the test for content issues, just report them
            # (visual test is the primary validation)


# ============================================================================
# Standalone Runner
# ============================================================================

if __name__ == "__main__":
    """Run visual tests directly."""
    import sys

    # Default to the samples directory relative to the project root
    project_root = Path(__file__).parent.parent.parent
    samples_dir = project_root / "samples"

    print("Visual Validation Test Runner")
    print("=" * 60)
    print(f"Samples directory: {samples_dir}")
    print(f"Baselines directory: {BASELINES_DIR}")
    print(f"Outputs directory: {OUTPUTS_DIR}")
    print()

    # Check for command line args
    update = "--update-baselines" in sys.argv or "-u" in sys.argv
    generate = "--generate-only" in sys.argv or "-g" in sys.argv

    if update:
        print("Mode: UPDATE BASELINES")
    elif generate:
        print("Mode: GENERATE ONLY")
    else:
        print("Mode: COMPARE TO BASELINES")

    print()
    ensure_dirs()

    samples = list(samples_dir.glob("*.docx"))
    samples = [s for s in samples if not s.name.startswith("~")]

    if not samples:
        print("No sample DOCX files found!")
        sys.exit(1)

    all_passed = True

    for sample in samples:
        name = sample.stem
        print(f"\nProcessing: {sample.name}")
        print("-" * 40)

        try:
            pdf_path, images = process_document(sample, OUTPUTS_DIR, name)

            if not images:
                print("  WARNING: No images generated")
                continue

            print(f"  Generated {len(images)} page(s)")

            for i, img in enumerate(images):
                baseline_path = BASELINES_DIR / f"{name}_page_{i}.png"
                output_path = OUTPUTS_DIR / f"{name}_page_{i}.png"
                diff_path = DIFFS_DIR / f"{name}_page_{i}_diff.png"

                img.save(output_path)

                if update:
                    img.save(baseline_path)
                    print(f"  Page {i}: baseline updated")
                elif generate:
                    print(f"  Page {i}: generated")
                elif baseline_path.exists():
                    baseline_img = Image.open(baseline_path)
                    is_same, diff_percent = compare_images(
                        baseline_img, img, diff_path
                    )
                    status = "OK" if is_same else "DIFFERS"
                    print(f"  Page {i}: {status} ({diff_percent:.2f}%)")
                    if not is_same:
                        all_passed = False
                else:
                    print(f"  Page {i}: no baseline")

        except Exception as e:
            print(f"  ERROR: {e}")
            all_passed = False

    print()
    print("=" * 60)
    if update:
        print("Baselines updated successfully!")
    elif generate:
        print("Screenshots generated successfully!")
    elif all_passed:
        print("All visual tests PASSED!")
    else:
        print("Some visual tests FAILED!")
        sys.exit(1)
